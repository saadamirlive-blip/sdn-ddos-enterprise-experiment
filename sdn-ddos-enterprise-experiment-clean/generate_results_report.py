"""
Results & Figures Report Generator
=====================================
Builds a complete Word document with every figure and table your paper
needs -- Traffic timeline, Availability timeline, Attack vector breakdown,
Severity/Response-time analysis, comparison bar charts, radar chart, and
summary tables -- ENTIRELY from real files produced by your actual runs:

  model_metrics.json                   <- from train_model.py (Step 1)
  trial_XX/real_metrics_summary.json   <- from metrics_from_logs.py (Step 4, per trial)
  trial_XX/ground_truth.json           <- from run_experiment.py
  logs/decisions_*.jsonl               <- from enterprise_security_controller_v2.py
  trial_XX/availability_*.jsonl        <- from availability_prober.py

NOTHING in this script is hardcoded. If a required input file is missing,
the corresponding section is SKIPPED with a clear note in the report,
rather than filled with a placeholder or invented number.

Optional: baselines.json -- if you separately obtain or measure comparison
figures for Traditional Firewall / IDS-IPS / Static SDN, provide them here.
This is NEVER auto-generated -- see the template printed by --init-baselines.
Without this file, comparison charts show ONLY your real, measured system.

Usage:
    python3 generate_results_report.py \\
        --model-metrics model_metrics.json \\
        --trials trial_01 trial_02 trial_03 \\
        --decisions-dir logs/ \\
        --baselines baselines.json \\
        --out Results_Report.docx

    # To see the optional baselines.json template:
    python3 generate_results_report.py --init-baselines
"""

import argparse
import glob
import json
import os
import statistics

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, RGBColor

FIGDIR = 'report_figures'
os.makedirs(FIGDIR, exist_ok=True)

COLORS = {
    'proposed': '#2E7D32', 'firewall': '#C62828', 'idsips': '#EF6C00',
    'staticsdn': '#1565C0', 'block': '#C62828', 'quarantine': '#EF6C00',
    'ratelimit': '#F9A825', 'monitor': '#2E7D32',
}

TARGETS = {
    'containment_rate_percent': 96.8,
    'response_time_seconds': 2.30,
    'false_positive_rate_percent': 1.2,
    'network_availability_percent': 98.5,
    'false_containment_rate_percent': 1.5,
}

BASELINES_TEMPLATE = {
    "_comment": "OPTIONAL. Fill this in only with numbers you actually "
                "measured or can cite from a real source. Do NOT copy "
                "figures from an old paper draft. Delete any paradigm you "
                "don't have real data for -- the report will simply omit it.",
    "Traditional Firewall": {"containment_rate_percent": None, "response_time_seconds": None,
                              "false_positive_rate_percent": None, "false_containment_rate_percent": None,
                              "network_availability_percent": None},
    "IDS/IPS": {"containment_rate_percent": None, "response_time_seconds": None,
                "false_positive_rate_percent": None, "false_containment_rate_percent": None,
                "network_availability_percent": None},
    "Static SDN": {"containment_rate_percent": None, "response_time_seconds": None,
                    "false_positive_rate_percent": None, "false_containment_rate_percent": None,
                    "network_availability_percent": None},
}


def load_json(path):
    if not path or not os.path.exists(path):
        return None
    with open(path) as f:
        return json.load(f)


def load_baselines(path):
    data = load_json(path)
    if not data:
        return None, None
    if 'metrics' in data:
        return data['metrics'], data.get('status', 'provided_reference_values')
    return data, 'provided_reference_values'


def load_reference_proposed(path):
    data = load_json(path)
    if not data:
        return None
    comparison = data.get('comparison_metrics', data)
    return comparison.get('proposed_sdn_ml_dynamic')


def load_jsonl(path):
    if not path or not os.path.exists(path):
        return []
    records = []
    with open(path) as f:
        for line in f:
            line = line.strip()
            if line:
                records.append(json.loads(line))
    return records


def mean_std(values):
    values = [v for v in values if v is not None]
    if not values:
        return None, None
    if len(values) == 1:
        return values[0], 0.0
    return statistics.mean(values), statistics.stdev(values)


# ----------------------------------------------------------------------- #
#  Figure builders -- each returns a PNG path, or None if data is missing  #
# ----------------------------------------------------------------------- #

def fig_traffic_timeline(decisions, out_path):
    if not decisions:
        return None
    t0 = min(d['t_detect_start'] for d in decisions)
    times = [(d['t_detect_start'] - t0) for d in decisions]
    pps = [d.get('pps', 0) for d in decisions]
    actions = [d['action'] for d in decisions]

    fig, ax = plt.subplots(figsize=(9, 4))
    ax.plot(times, pps, color='#555', linewidth=1, alpha=0.6, label='Observed pps')
    for act, color in [('BLOCK', COLORS['block']), ('QUARANTINE', COLORS['quarantine']),
                        ('RATE_LIMIT', COLORS['ratelimit'])]:
        xs = [t for t, a in zip(times, actions) if a == act]
        ys = [p for p, a in zip(pps, actions) if a == act]
        if xs:
            ax.scatter(xs, ys, color=color, label=f'{act} triggered', zorder=5, s=20)
    ax.set_xlabel('Time (seconds since trial start)')
    ax.set_ylabel('Packets per second (pps)')
    ax.set_title('Traffic Volume Over Time (Real Trial Data)')
    ax.legend(fontsize=8)
    ax.grid(alpha=0.3)
    fig.tight_layout()
    fig.savefig(out_path, dpi=150)
    plt.close(fig)
    return out_path


def fig_availability_timeline(avail_records, out_path):
    if not avail_records:
        return None
    t0 = min(r['timestamp'] for r in avail_records)
    bucket_s = 5
    buckets = {}
    for r in avail_records:
        b = int((r['timestamp'] - t0) // bucket_s)
        buckets.setdefault(b, []).append(r['success'])
    xs = sorted(buckets.keys())
    ys = [100 * sum(buckets[b]) / len(buckets[b]) for b in xs]
    times = [x * bucket_s for x in xs]

    fig, ax = plt.subplots(figsize=(9, 4))
    ax.plot(times, ys, color=COLORS['proposed'], linewidth=1.5)
    ax.axhline(95, color='red', linestyle='--', linewidth=1, label='95% SLA threshold')
    ax.set_xlabel('Time (seconds since trial start)')
    ax.set_ylabel('Network Availability (%)')
    ax.set_ylim(0, 105)
    ax.set_title('Network Availability Over Time (Real Trial Data)')
    ax.legend(fontsize=8)
    ax.grid(alpha=0.3)
    fig.tight_layout()
    fig.savefig(out_path, dpi=150)
    plt.close(fig)
    return out_path


def fig_attack_breakdown(per_attack_stats, out_path):
    if not per_attack_stats:
        return None
    types = list(per_attack_stats.keys())
    rates = [per_attack_stats[t]['containment_rate_percent'] for t in types]
    fig, ax = plt.subplots(figsize=(8, 4.5))
    bars = ax.bar(types, rates, color=COLORS['proposed'])
    for b, r in zip(bars, rates):
        if r is not None:
            ax.text(b.get_x() + b.get_width()/2, r + 1, f'{r:.1f}%', ha='center', fontsize=9)
    ax.set_ylabel('Containment Rate (%)')
    ax.set_title('Containment Rate by Attack Type (Real, Aggregated Across Trials)')
    ax.set_ylim(0, 110)
    plt.setp(ax.get_xticklabels(), rotation=30, ha='right')
    fig.tight_layout()
    fig.savefig(out_path, dpi=150)
    plt.close(fig)
    return out_path


def fig_severity_response(all_decisions, out_path):
    contained = [d for d in all_decisions if d['action'] in ('RATE_LIMIT', 'QUARANTINE', 'BLOCK')]
    if not contained:
        return None
    fig, axes = plt.subplots(1, 2, figsize=(11, 4.5))
    for act, color in [('MONITOR', COLORS['monitor']), ('RATE_LIMIT', COLORS['ratelimit']),
                        ('QUARANTINE', COLORS['quarantine']), ('BLOCK', COLORS['block'])]:
        sev = [d['severity'] for d in all_decisions if d['action'] == act]
        rt = [d['total_s'] for d in all_decisions if d['action'] == act]
        if sev:
            axes[0].scatter(sev, rt, color=color, label=act, alpha=0.5, s=12)
    axes[0].set_xlabel('Severity (P_attack)')
    axes[0].set_ylabel('Response Time (s)')
    axes[0].set_title('Severity vs. Response Time')
    axes[0].legend(fontsize=8)
    axes[0].grid(alpha=0.3)

    box_data, labels = [], []
    for act in ['MONITOR', 'RATE_LIMIT', 'QUARANTINE', 'BLOCK']:
        rt = [d['total_s'] for d in all_decisions if d['action'] == act]
        if rt:
            box_data.append(rt)
            labels.append(act)
    if box_data:
        try:
            axes[1].boxplot(box_data, tick_labels=labels)
        except TypeError:
            axes[1].boxplot(box_data, labels=labels)  # older matplotlib
        axes[1].set_ylabel('Response Time (s)')
        axes[1].set_title('Response Time Distribution by Action')
        plt.setp(axes[1].get_xticklabels(), rotation=20)
    fig.tight_layout()
    fig.savefig(out_path, dpi=150)
    plt.close(fig)
    return out_path


def fig_comparison_bar(metric_key, ylabel, title, proposed_value, baselines, out_path):
    labels, values, colors = ['Proposed SDN\n(Real)'], [proposed_value], [COLORS['proposed']]
    key_map = {'Traditional Firewall': 'firewall', 'IDS/IPS': 'idsips', 'Static SDN': 'staticsdn'}
    for name, data in (baselines or {}).items():
        if name.startswith('_'):
            continue
        v = data.get(metric_key)
        if v is not None:
            labels.append(f'{name}\n(external ref.)')
            values.append(v)
            colors.append(COLORS.get(key_map.get(name, ''), '#888'))
    note = "  (no baseline data provided)" if len(values) == 1 else ""
    fig, ax = plt.subplots(figsize=(7, 4.5))
    bars = ax.bar(labels, values, color=colors)
    for b, v in zip(bars, values):
        ax.text(b.get_x() + b.get_width()/2, v, f'{v:.2f}', ha='center', va='bottom', fontsize=9)
    ax.set_ylabel(ylabel)
    ax.set_title(title + note, fontsize=10)
    fig.tight_layout()
    fig.savefig(out_path, dpi=150)
    plt.close(fig)
    return out_path


def fig_radar(proposed, baselines, out_path):
    metrics = ['containment_rate_percent', 'network_availability_percent']
    inv_metrics = ['response_time_seconds', 'false_positive_rate_percent', 'false_containment_rate_percent']
    labels = ['Containment\nRate', 'Availability', 'Speed\n(inv. resp. time)', 'Low FPR', 'Low FCR']

    def to_radar_values(d):
        vals = []
        for m in metrics:
            vals.append(d.get(m) if d.get(m) is not None else 0)
        for m in inv_metrics:
            v = d.get(m)
            vals.append(100 - v if v is not None else 0)
        return vals

    series = {'Proposed SDN (Real)': to_radar_values(proposed)}
    for name, data in (baselines or {}).items():
        if name.startswith('_'):
            continue
        if any(v is not None for v in data.values()):
            series[f'{name} (external ref.)'] = to_radar_values(data)

    angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
    for name, vals in series.items():
        vals = vals + vals[:1]
        ax.plot(angles, vals, label=name, linewidth=1.5)
        ax.fill(angles, vals, alpha=0.08)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels, fontsize=8)
    ax.set_title('Multi-Metric Performance Radar (Real Data)', fontsize=10)
    ax.legend(loc='upper right', bbox_to_anchor=(1.3, 1.1), fontsize=7)
    fig.tight_layout()
    fig.savefig(out_path, dpi=150)
    plt.close(fig)
    return out_path


# ----------------------------------------------------------------------- #
#  Document assembly                                                      #
# ----------------------------------------------------------------------- #

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_note(doc, text, missing=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xB0, 0x00, 0x20) if missing else RGBColor(0x66, 0x66, 0x66)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = str(h)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--model-metrics', default='model_metrics.json')
    ap.add_argument('--trials', nargs='*', default=[], help='List of trial directories, e.g. trial_01 trial_02 ...')
    ap.add_argument('--decisions-dir', default='logs', help='Directory containing decisions_*.jsonl from the controller')
    ap.add_argument('--baselines', default=None, help='Optional baselines.json path')
    ap.add_argument('--reference-metrics', default=None,
                    help='Optional supplied paper metrics JSON; shown separately from measured trials')
    ap.add_argument('--out', default='Results_Report.docx')
    ap.add_argument('--init-baselines', action='store_true', help='Print the optional baselines.json template and exit')
    args = ap.parse_args()

    if args.init_baselines:
        print(json.dumps(BASELINES_TEMPLATE, indent=2))
        return

    doc = Document()
    doc.add_heading('Experimental Results Report', level=0)
    p = doc.add_paragraph()
    p.add_run('Generated entirely from real experiment logs -- no invented or placeholder values.').italic = True

    missing_sections = []

    add_heading(doc, '1. Machine Learning Classification Performance', 1)
    mm = load_json(args.model_metrics)
    if mm:
        confusion = mm.get('confusion_matrix', {})
        add_table(doc, ['Metric', 'Value'], [
            ['Test samples', mm.get('n_test_samples', confusion.get('n_test_samples'))],
            ['Accuracy', f"{mm['accuracy_percent']}%"],
            ['Precision', f"{mm['precision_percent']}%"],
            ['Recall', f"{mm['recall_percent']}%"],
            ['ROC-AUC', mm['roc_auc']],
            ['True Normal (TN)', confusion.get('true_normal', confusion.get('true_negative'))],
            ['False Positive (FP)', confusion.get('false_positive')],
            ['False Negative (FN)', confusion.get('false_negative')],
            ['True Attack (TP)', confusion.get('true_attack', confusion.get('true_positive'))],
        ])
        doc.add_heading('Feature Importance', level=2)
        add_table(doc, ['Feature', 'Importance'],
                  [[k, v] for k, v in mm['feature_importance'].items()])
    else:
        add_note(doc, 'SKIPPED: model_metrics.json not found. Run train_model.py (with the JSON-export patch) first.', missing=True)
        missing_sections.append('ML classification performance')

    trial_summaries = []
    all_decisions = []
    for t in args.trials:
        s = load_json(os.path.join(t, 'real_metrics_summary.json'))
        if s:
            trial_summaries.append(s)
    trial_decision_files = [os.path.join(t, 'decisions.jsonl') for t in args.trials
                            if os.path.exists(os.path.join(t, 'decisions.jsonl'))]
    if trial_decision_files:
        for df in trial_decision_files:
            all_decisions.extend(load_jsonl(df))
    else:
        dec_files = glob.glob(os.path.join(args.decisions_dir, 'decisions_*.jsonl'))
        for df in dec_files:
            all_decisions.extend(load_jsonl(df))

    proposed = {}
    add_heading(doc, '2. Aggregated Performance Across Trials', 1)
    if trial_summaries:
        cr_mean, cr_std = mean_std([s['detection_and_containment']['containment_rate_percent'] for s in trial_summaries])
        fpr_mean, fpr_std = mean_std([s['detection_and_containment']['false_positive_rate_percent'] for s in trial_summaries])
        fcr_mean, fcr_std = mean_std([s['detection_and_containment'].get('false_containment_rate_percent') for s in trial_summaries])
        rt_mean, rt_std = mean_std([s['detection_and_containment']['response_time_seconds']['mean'] for s in trial_summaries])
        na_mean, na_std = mean_std([s['availability']['network_availability_percent'] if s.get('availability') else None for s in trial_summaries])

        def fmt(v):
            return f'{v:.2f}' if v is not None else 'N/A'

        add_table(doc, ['Metric', 'Mean', 'Std Dev', 'N Trials'], [
            ['Containment Rate (%)', fmt(cr_mean), fmt(cr_std), len(trial_summaries)],
            ['False Positive Rate (%)', fmt(fpr_mean), fmt(fpr_std), len(trial_summaries)],
            ['False Containment Rate (%)', fmt(fcr_mean), fmt(fcr_std), len(trial_summaries)],
            ['Response Time (s)', fmt(rt_mean), fmt(rt_std), len(trial_summaries)],
            ['Network Availability (%)', fmt(na_mean), fmt(na_std), len(trial_summaries)],
        ])
        proposed = {
            'containment_rate_percent': cr_mean, 'false_positive_rate_percent': fpr_mean,
            'false_containment_rate_percent': fcr_mean,
            'response_time_seconds': rt_mean, 'network_availability_percent': na_mean,
        }
    else:
        add_note(doc, 'SKIPPED: no trial real_metrics_summary.json files found. Run Steps 3-5 (run_experiment.py + metrics_from_logs.py) first.', missing=True)
        missing_sections.append('Aggregated trial performance')

    if proposed:
        add_heading(doc, 'Target Comparison', 2)
        target_rows = []
        for key, label in [
            ('containment_rate_percent', 'Containment Rate (%)'),
            ('response_time_seconds', 'Response Time (s)'),
            ('false_positive_rate_percent', 'False Positive Rate (%)'),
            ('network_availability_percent', 'Network Availability (%)'),
            ('false_containment_rate_percent', 'False Containment Rate (%)'),
        ]:
            measured = proposed.get(key)
            target = TARGETS[key]
            delta = measured - target if measured is not None else None
            target_rows.append([label, fmt(measured), f'{target:.2f}', fmt(delta)])
        add_table(doc, ['Metric', 'Measured Mean', 'Target', 'Difference'], target_rows)

    reference_proposed = load_reference_proposed(args.reference_metrics)
    if reference_proposed:
        add_heading(doc, 'Supplied Paper Reference Results', 2)
        add_note(doc, f'Reference values loaded from {args.reference_metrics}; these are supplied summary values, not rerun measurements.')
        reference_rows = []
        for key, label in [
            ('containment_rate_percent', 'Containment Rate (%)'),
            ('response_time_seconds', 'Response Time (s)'),
            ('false_positive_rate_percent', 'False Positive Rate (%)'),
            ('network_availability_percent', 'Network Availability (%)'),
            ('false_containment_rate_percent', 'False Containment Rate (%)'),
        ]:
            reference_rows.append([label, reference_proposed.get(key)])
        add_table(doc, ['Metric', 'Supplied Reference'], reference_rows)

    baselines, baseline_status = load_baselines(args.baselines)
    if args.baselines and not baselines:
        add_note(doc, f'NOTE: --baselines {args.baselines} was specified but not found/empty -- comparison charts show only the proposed system.', missing=True)
    elif baselines and baseline_status:
        add_note(doc, f'Baseline values supplied externally ({baseline_status}); they were not rerun by this pipeline.')

    add_heading(doc, '3. Traffic Volume Over Time', 1)
    if all_decisions:
        rep = fig_traffic_timeline(sorted(all_decisions, key=lambda d: d['t_detect_start'])[:2000],
                                    os.path.join(FIGDIR, 'traffic_timeline.png'))
        if rep:
            doc.add_picture(rep, width=Inches(6))
    else:
        add_note(doc, 'SKIPPED: no decisions log data found.', missing=True)
        missing_sections.append('Traffic timeline')

    add_heading(doc, '4. Network Availability Over Time', 1)
    avail_recs = []
    for t in args.trials:
        for af in glob.glob(os.path.join(t, 'availability_*.jsonl')):
            avail_recs = load_jsonl(af)
            if avail_recs:
                break
        if avail_recs:
            break
    if avail_recs:
        f = fig_availability_timeline(avail_recs, os.path.join(FIGDIR, 'availability_timeline.png'))
        if f:
            doc.add_picture(f, width=Inches(6))
    else:
        add_note(doc, 'SKIPPED: no availability_*.jsonl found in any trial directory.', missing=True)
        missing_sections.append('Availability timeline')

    add_heading(doc, '5. Containment Efficacy by Attack Vector', 1)
    per_attack = {}
    for t in args.trials:
        gt = load_json(os.path.join(t, 'ground_truth.json'))
        if not gt:
            continue
        for w in gt['attack_windows']:
            atype = w['attack_type']
            per_attack.setdefault(atype, {'total': 0, 'contained': 0})
            per_attack[atype]['total'] += 1
            hit = any(w['start'] <= d['t_detect_start'] <= w['end'] and d['src_ip'] == w['src_ip']
                      and d['action'] in ('RATE_LIMIT', 'QUARANTINE', 'BLOCK')
                      for d in all_decisions)
            if hit:
                per_attack[atype]['contained'] += 1
    per_attack_stats = {
        k: {'containment_rate_percent': round(100 * v['contained'] / v['total'], 1) if v['total'] else None}
        for k, v in per_attack.items()
    }
    if per_attack_stats:
        add_table(doc, ['Attack Type', 'Containment Rate (%)'],
                  [[k, v['containment_rate_percent']] for k, v in per_attack_stats.items()])
        f = fig_attack_breakdown(per_attack_stats, os.path.join(FIGDIR, 'attack_breakdown.png'))
        if f:
            doc.add_picture(f, width=Inches(6))
    else:
        add_note(doc, 'SKIPPED: no ground_truth.json attack windows found across provided trials.', missing=True)
        missing_sections.append('Attack vector breakdown')

    add_heading(doc, '6. Enforcement Severity vs. Response Time', 1)
    f = fig_severity_response(all_decisions, os.path.join(FIGDIR, 'severity_response.png'))
    if f:
        doc.add_picture(f, width=Inches(6.5))
    else:
        add_note(doc, 'SKIPPED: no contained (RATE_LIMIT/QUARANTINE/BLOCK) decisions found in logs.', missing=True)
        missing_sections.append('Severity vs response time')

    add_heading(doc, '7. Comparative Benchmarks', 1)
    if proposed:
        specs = [
            ('containment_rate_percent', 'Containment Rate (%)', 'Threat Containment Rate'),
            ('response_time_seconds', 'Response Time (s)', 'Average Response Latency'),
            ('false_positive_rate_percent', 'FPR (%)', 'False Positive Rate'),
            ('false_containment_rate_percent', 'FCR (%)', 'False Containment Rate'),
            ('network_availability_percent', 'Availability (%)', 'Network Service Availability'),
        ]
        any_chart = False
        for key, ylabel, title in specs:
            val = proposed.get(key)
            if val is None:
                continue
            f = fig_comparison_bar(key, ylabel, title, val, baselines,
                                    os.path.join(FIGDIR, f'compare_{key}.png'))
            doc.add_picture(f, width=Inches(5.5))
            any_chart = True
        if any_chart:
            f = fig_radar(proposed, baselines, os.path.join(FIGDIR, 'radar.png'))
            doc.add_picture(f, width=Inches(5))
    else:
        add_note(doc, 'SKIPPED: no aggregated proposed-system metrics available.', missing=True)
        missing_sections.append('Comparative benchmarks')

    add_heading(doc, '8. Summary', 1)
    if missing_sections:
        add_note(doc, 'The following sections were skipped due to missing input data: ' +
                  '; '.join(missing_sections) + '. Re-run this script once those files exist.', missing=True)
    else:
        doc.add_paragraph('All sections generated successfully from real experiment data.')

    doc.save(args.out)
    print(f"\nSaved: {args.out}")
    if missing_sections:
        print(f"WARNING: {len(missing_sections)} section(s) were skipped -- see the report for details.")


if __name__ == '__main__':
    main()
